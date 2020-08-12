from docx import Document

from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

class MyTools_word(object):
    def __init__(self):
        p = document.add_paragraph('')
        return p
    def new_word():
        document = Document()
        
    def add_paragraph(type='paragraph'):
        if type=='paragraph':
            p = document.add_paragraph('')
        elif type=='heading':
            p = document.add_heading('')
        else:
            pirnt('type类型错误')
            assert 0
        return p

    def save_word(paths_name='c:/Users/Administrator/Desktop/myword.docx'):
        document.save(paths_name)