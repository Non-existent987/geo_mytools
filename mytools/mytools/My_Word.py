#!/usr/bin/env python
# coding: utf-8

from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

class MyTools_word(object):
    document = Document()
    p = document.add_paragraph('')
    def add_heading(self,
                    document,
                    wenzi='MR竞对覆盖评估体系',
                    jibie=1,size=11,
                    ziti = u'微软雅黑',
                    yanse = RGBColor(0,0,0),
                    duiqi='l',
                    bold_is=True):
        '''
        jibie = 1-9
        size = 5-72
        ziti = u'宋体' or u'微软雅黑'
        yanse = RGBColor(red,green,blue)(255,255,255)
        duiqi = c(CENTER) or l(LEFT) or r(RIGHT)
        bold_is=True
        
        示例：
        import mytools
        from docx import Document
        import pandas as pd
        w = mytools.wordn
        cmcc_res = pd.read_csv('G:/cmcc_res.csv',encoding='gbk')
        document = Document()
        w.add_heading(document,'标题1')
        w.add_paragraph(document,'段落1')
        p = document.add_paragraph('')
        w.add_text(p,'添加的信息1')
        w.add_text(p,'添加的信息2')
        w.add_paragraph(document,'段落2')
        w.add_df(document,cmcc_res,'g',lists=[((1,1),(3,2)),((4,2),(6,3))])
        document.save('demo.docx')
        '''

        biaoti = document.add_heading('',jibie)
        if duiqi == 'c':
            biaoti.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif duiqi == 'l':
            biaoti.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            biaoti.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        tt=biaoti.add_run(wenzi)
        tt.font.size = Pt(size)
        if bold_is:
            tt.bold
        else:
            pass
        tt.font.name=ziti
        tt._element.rPr.rFonts.set(qn('w:eastAsia'), ziti) 
        tt.font.color.rgb = RGBColor(0,0,0)
        paragraph_format = biaoti.paragraph_format
        paragraph_format.line_spacing = 1.5 # 1.5倍行间距
    def add_text(self,
                    p,
                    wenzi='文字描述',
                    size=11,
                    bold=False ,
                    ziti =u'宋体',
                    color='m'):
        '''
        size = 5-72
        ziti = u'宋体' or u'微软雅黑'
        bold=True

        示例：
        import mytools
        from docx import Document
        import pandas as pd
        w = mytools.wordn
        cmcc_res = pd.read_csv('G:/cmcc_res.csv',encoding='gbk')
        document = Document()
        w.add_heading(document,'标题1')
        w.add_paragraph(document,'段落1')
        p = document.add_paragraph('')
        w.add_text(p,'添加的信息1')
        w.add_text(p,'添加的信息2')
        w.add_paragraph(document,'段落2')
        w.add_df(document,cmcc_res,'g',lists=[((1,1),(3,2)),((4,2),(6,3))])
        document.save('demo.docx')
        '''
        
        run1 = p.add_run(wenzi)  # 使用add_run添加文字
        run1.font.size = Pt(size) #字体大小设置，和word里面的字号相对应，小一
        if bold:
            run1.font.bold = True
        run1.font.name=ziti
        if color=='m':
            run1.font.color.rgb = RGBColor(0, 0, 0)
        elif color=='r':
            run1.font.color.rgb = RGBColor(255, 0, 0)
        elif color=='b':
            run1.font.color.rgb = RGBColor(0, 0, 255)
        else:
            run1.font.color.rgb = RGBColor(0, 0, 0)
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), ziti) 
        paragraph_format = p.paragraph_format
        paragraph_format.first_line_indent = Inches(0.25)
        paragraph_format.line_spacing = 1.5 # 1.5倍行间距
    def add_paragraph(self,
                        document,
                        wenzi='文字描述',
                        size=11,
                        bold=False ,
                        ziti =u'宋体'
                        ):
        '''
        size = 5-72
        ziti = u'宋体' or u'微软雅黑'
        bold=True

        示例：
        import mytools
        from docx import Document
        import pandas as pd
        w = mytools.wordn
        cmcc_res = pd.read_csv('G:/cmcc_res.csv',encoding='gbk')
        document = Document()
        w.add_heading(document,'标题1')
        w.add_paragraph(document,'段落1')
        p = document.add_paragraph('')
        w.add_text(p,'添加的信息1')
        w.add_text(p,'添加的信息2')
        w.add_paragraph(document,'段落2')
        w.add_df(document,cmcc_res,'g',lists=[((1,1),(3,2)),((4,2),(6,3))])
        document.save('demo.docx')
        '''
        
        p = document.add_paragraph('')
        run1 = p.add_run(wenzi)  # 使用add_run添加文字
        run1.font.size = Pt(size) #字体大小设置，和word里面的字号相对应，小一
        if bold:
            run1.font.bold = True
        run1.font.name=ziti
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), ziti) 
        paragraph_format = p.paragraph_format
        paragraph_format.first_line_indent = Inches(0.25)
        paragraph_format.line_spacing = 1.5 # 1.5倍行间距
    def add_df(self,
                document,
                df,
                color='m',
                size=11,
                lists=[((0,0),(0,0))],
                bold=True,
                lie_k = False,
                lie_k_lie=[10,4],
                zishiying=True,
                int_0=False
                ):
        '''
        document = Document()
        df = DataFrame()
        color=m(默认) or r(红色) or b(蓝色) or g(绿色)
        lists=[((起始单元格的行,起始单元格的列),(结束单元格的行,结束单元格的列))]，默认从0开始
        size = 5-72
        ziti = u'宋体' or u'微软雅黑'
        bold=True

        示例：
        import mytools
        from docx import Document
        import pandas as pd
        w = mytools.wordn
        cmcc_res = pd.read_csv('G:/cmcc_res.csv',encoding='gbk')
        document = Document()
        w.add_heading(document,'标题1')
        w.add_paragraph(document,'段落1')
        p = document.add_paragraph('')
        w.add_text(p,'添加的信息1')
        w.add_text(p,'添加的信息2')
        w.add_paragraph(document,'段落2')
        w.add_df(document,cmcc_res,'g',lists=[((1,1),(3,2)),((4,2),(6,3))])
        document.save('demo.docx')
        '''
        
        #插入表格并赋值
        data = df.copy()
        document.styles['Table Grid'].font.name = u'微软雅黑'
        document.styles['Table Grid']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        
        table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
        if zishiying:
            table.autofit=True
        else:
            table.autofit=False
        table.rows[0].height=Cm(1)
         #合并单元格
        for list_t in lists:
            x,y=list_t[0]
            x1,y1=list_t[1]
            cell_1=table.cell(x, y)
            cell_2=table.cell(x1, y1)
            cell_3 = cell_1.merge(cell_2)
            cell_3.txt = data.iloc[x, y]
        # 列名赋值
        hdr_cells = table.rows[0].cells
        for row in list(data.columns):
            for num in range(0, data.shape[1]):
                # hdr_cells[num].text = u'' + list(data.columns)[num]
                # table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
                # table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                #添加内容并修改样式
                run=table.cell(0,num).paragraphs[0].add_run(u'' + list(data.columns)[num])
                if bold:
                    run.font.bold=True
                run.font.size= Pt(size)
                table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
                table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if lie_k:
                    lie_k_ji,lie_k_n = lie_k_lie
                    if num==lie_k_ji:          ##设置第一列的列宽
                        table.cell(0,num).width=Inches(lie_k_n)
                        # table.cell(row,col).width=Inches()
                # run.font.name='宋体'
                # run.font.size=Pt(size)
                # run.font.color.rgb=RGBColor(233,123,12)
                

                if color=='m':
                    mycolor='#BFBFBF'
                elif color=='r':
                    mycolor='#F79646'
                elif color=='b':
                    mycolor='#4F81BD'
                elif color=='g':
                    mycolor='#4BACC6'
                else:
                    mycolor='#BFBFBF'
                shading_elm_2 = parse_xml(r'<w:shd {} w:fill="{name_b}"/>'.format(nsdecls('w'),name_b=mycolor))
                table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
            break
        # 内容赋值
        for row in range(1,data.shape[0]+1):
            row_cells = table.rows[row].cells
            for num in range(0, data.shape[1]):
                #添加内容并修改样式
                if int_0:
                    try:
                        data_cell_str = str(int(data.iloc[row-1,num]))
                    except:
                        data_cell_str = str(data.iloc[row-1,num])
                else:
                    data_cell_str = str(data.iloc[row-1,num])

                run2=table.cell(row,num).paragraphs[0].add_run(u'' + data_cell_str)
                run2.font.size = Pt(size)
                # row_cells[num].text = u'' + str(data.iloc[row-1,num])
                table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
                table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # #单个单元格设置
        # run.font.color.rgb = RGBColor(255, 0, 0)  # 颜色设置，这里是用RGB颜色
        # run.font.size = Pt(15)  # 字体大小设置，和word里面的字号相对应
        # p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    #设置对齐方式

        #设置整个表格字体属性
        # table.style.font.size=Pt(size)
        # table.style.font.bold=True
        table.style.font.color.rgb=RGBColor(0, 0, 0)
        table.style.paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
# a = MyTools_word()