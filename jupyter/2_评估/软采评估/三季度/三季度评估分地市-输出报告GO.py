# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL

#评估标准插入
biaozhun = pd.read_csv(open('d:/shiyong/评估标准_插入.csv',encoding = 'gbk'))

#软采小区替换竞对后数据
ruancai = pd.read_csv(open('d:/shiyong/汇总-添评估使用-备份-华为使用邮件数据-tdd-不要大唐.csv',encoding = 'gbk'))

#软采栅格
data_merge_quchong1 = pd.read_csv(open('d:/shiyong/第三季度_栅格_添加场景_添加区县_添加弱覆盖1.csv',encoding = 'gbk'))

#对应关系，县城和地市
data_duiyingguanxi = pd.read_csv(open('d:/shiyong/县城和地市对应关系1.csv',encoding = 'gbk'))

data_duiyingguanxi.columns = ['quxian_xxx','dishi_xxx']

data_1_rres = data_merge_quchong1.merge(data_duiyingguanxi,how='left',left_on='区县',right_on='quxian_xxx')

data_merge_quchong = data_1_rres[['a.grid_id', 'a.day', 'a.grid_leftbottom_latitude',
       'a.grid_leftbottom_longitude', 'a.grid_lefttop_latitude',
       'a.grid_lefttop_longitude', 'a.grid_rightbottom_latitude',
       'a.grid_rightbottom_longitude', 'a.grid_righttop_latitude',
       'a.grid_righttop_longitude', 'a.inter_maxrsrp_weak', 'a.region_name',
       'a.rsrp_sample_count', 'a.rsrp_sum', 'a.rsrp_weak', 'a.scrsrp_110',
       'a.scrsrp_110_95', 'a.scrsrp_95', 'a.te_maxrsrp_sample_count',
       'a.te_maxrsrp_sum', 'a.te_maxrsrp_weak', 'a.te_maxrsrp_weak_l2',
       'a.un_maxrsrp_sample_count', 'a.un_maxrsrp_sum', 'a.un_maxrsrp_weak',
       'a.un_maxrsrp_weak_l2', 'b.dl_data', 'b.ul_data', 'lon', 'lat', '区域类型',
       '区县', '全网', 'fgl', '大于100采样点', '弱覆盖栅格', '严重弱覆盖栅格','dishi_xxx']]
list_dishi = ['孝感','襄阳','咸宁','随州','十堰','荆州','荆门','黄石','鄂州','黄冈','恩施','宜昌','潜江','天门','江汉']


def to_word(dishi_name = 'name'):
    dishi =['十堰','荆门','武汉','襄阳','黄石','荆州','宜昌','鄂州','孝感','潜江','江汉','黄冈','天门','随州','恩施','咸宁','武汉']
    document = Document()
    def style_heading(wenzi = 'MR竞对覆盖评估体系',jibie = 1,size = 16,ziti = u'微软雅黑',yanse = RGBColor(0,0,0),duiqi = 'CENTER'):
        biaoti = document.add_heading('',jibie)
        if duiqi == 'CENTER':
            biaoti.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif duiqi == 'LEFT':
            biaoti.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            biaoti.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        tt=biaoti.add_run(wenzi)
        tt.font.size = Pt(size)
        tt.bold
        tt.font.name=ziti
        tt._element.rPr.rFonts.set(qn('w:eastAsia'), ziti) 
        tt.font.color.rgb = RGBColor(0,0,0)
        paragraph_format = biaoti.paragraph_format
        paragraph_format.line_spacing = 1.5 # 1.5倍行间距
    def style_paragraph(wenzi = '文字描述',size = 10.5,bold = False ,ziti =u'宋体' ):
        p = document.add_paragraph('')
        run1 = p.add_run(wenzi)  # 使用add_run添加文字
        run1.font.size = Pt(size) #字体大小设置，和word里面的字号相对应，小一
        if bold:
            run1.font.bold = True
        run1.font.name=ziti
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), ziti) 
        paragraph_format = p.paragraph_format
        paragraph_format.first_line_indent = Inches(0.25)
        paragraph_format.line_spacing = 1.5 # 1.5倍行间距
    def style_paragraph_add(wenzi = '文字描述',size = 10.5,bold = False ,ziti =u'宋体' ):
        run1 = p.add_run(wenzi)  # 使用add_run添加文字
        run1.font.size = Pt(size) #字体大小设置，和word里面的字号相对应，小一
        if bold:
            run1.font.bold = True
        run1.font.name=ziti
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), ziti) 
        paragraph_format = p.paragraph_format
        paragraph_format.first_line_indent = Inches(0.25)
        paragraph_format.line_spacing = 1.5 # 1.5倍行间距
    
    def style_paragraph_hongse_add(wenzi = '文字描述',size = 10.5,bold = False ,ziti =u'宋体' ):
        run1 = p.add_run(wenzi)  # 使用add_run添加文字
        run1.font.size = Pt(size) #字体大小设置，和word里面的字号相对应，小一
        font_color = run1.font
        font_color.color.rgb = RGBColor(255, 0 , 0)
        if bold:
            run1.font.bold = True
        run1.font.name=ziti
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), ziti) 
        paragraph_format = p.paragraph_format
        paragraph_format.first_line_indent = Inches(0.25)
        paragraph_format.line_spacing = 1.5 # 1.5倍行间距
    ruancai['全网']='全网'
    def aggfunc(df):
        sum1 = df['rsrp_sample_count'].sum()
        sum2 = df['rsrp_weak'].sum()
        ser = pd.Series([sum1, sum2, 1-sum2/sum1])
        return ser
    data_fgl = ruancai.groupby(by='region_name').apply(aggfunc)
    data_fgl.reset_index(inplace=True)
    data_fgl.columns = ['全网-指标','移动总点','移动弱点','全网移动覆盖率']
    data_fgl.sort_values(by = '全网移动覆盖率',ascending = False,inplace = True)
    #添加武汉中兴华为
    data_fg_changjia = ruancai.loc[ruancai['region_name']=='武汉'].groupby(by='厂家名称').apply(aggfunc)
    data_fg_changjia.reset_index(inplace=True)
    data_fg_changjia.columns = ['全网-指标','移动总点','移动弱点','全网移动覆盖率']
    data_fg_changjia2 = data_fg_changjia.loc[(data_fg_changjia['全网-指标']=='中兴') | (data_fg_changjia['全网-指标']=='华为')]
    data_fg_changjia2.loc[data_fg_changjia2['全网-指标']=='中兴','全网-指标']='武汉中兴'
    data_fg_changjia2.loc[data_fg_changjia2['全网-指标']=='华为','全网-指标']='武汉华为'
    data_fgl = pd.concat([data_fgl,data_fg_changjia2])
    #添加全网指标
    data_fg_quanwang = ruancai.groupby(by='全网').apply(aggfunc)
    data_fg_quanwang.reset_index(inplace=True)
    data_fg_quanwang.columns = ['全网-指标','移动总点','移动弱点','全网移动覆盖率']
    data_fgl = pd.concat([data_fgl,data_fg_quanwang])
    data_fgl['全网移动覆盖率'] = data_fgl['全网移动覆盖率'].apply(lambda x: format(x, '.2%'))    
    dishi_fgl_qian3=list(data_fgl.iloc[0:3,0])
    dishi_fgl_hou3=list(data_fgl.iloc[-6:-3,0])
    dishi_fgl_hou3.reverse()
    #各区县指标
    data_fgl_qx = ruancai.loc[ruancai['对应栅格地']==dishi_name].groupby(by='对应栅格区县').apply(aggfunc)
    data_fgl_qx.reset_index(inplace=True)
    data_fgl_qx.columns = ['区域','移动总点','移动弱点','全网移动覆盖率']
    data_fgl_qx.sort_values(by = '全网移动覆盖率',ascending = False,inplace = True)
    data_fgl_qx['全网移动覆盖率'] = data_fgl_qx['全网移动覆盖率'].apply(lambda x: format(x, '.2%'))   
    #dishi_name指标
    data_fgl_wuhan = data_fgl.loc[data_fgl['全网-指标']==dishi_name]
    data_fgl_wuhan.columns = ['区域','移动总点','移动弱点','全网移动覆盖率']
    #区县总计
    data_fgl_qx_ok = pd.concat([data_fgl_qx,data_fgl_wuhan])

    if data_fgl_qx_ok.shape[0]>=7:
        quxian_fgl_qian3=list(data_fgl_qx_ok.iloc[0:3,0])
        quxian_fgl_hou3=list(data_fgl_qx_ok.iloc[-4:-1,0])
        quxian_fgl_hou3.reverse()
    elif data_fgl_qx_ok.shape[0]>=5:
        quxian_fgl_qian3=list(data_fgl_qx_ok.iloc[0:2,0])
        quxian_fgl_hou3=list(data_fgl_qx_ok.iloc[-3:-1,0])
        quxian_fgl_hou3.reverse()
    else:
        quxian_fgl_qian3=list(data_fgl_qx_ok.iloc[0:1,0])
        quxian_fgl_hou3=list(data_fgl_qx_ok.iloc[-2:-1,0])
        quxian_fgl_hou3.reverse()
    
    #添加word内容
    style_heading()
    style_heading(wenzi ='一、	网络质量覆盖评估',jibie = 2,size = 14 , ziti=u'宋体',duiqi = 'LEFT')
    style_paragraph(wenzi = '此用于评估全网网络质量情况，并通过整网、小区、大数据栅格三个维度进行评估分析。')
    style_heading(wenzi ='1.整网粒度',jibie = 2,size = 12 , ziti=u'宋体',duiqi = 'LEFT')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='全网MR覆盖率：',bold = True)
    style_paragraph_add(wenzi='全省分地市，武汉分区县整体MR覆盖率情况。')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='指标定义：',bold = True)
    style_paragraph_add(wenzi='全网MR覆盖率=所有小区MR RSRP≥-110dbm的采样点总数/采样点总数')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='数据来源：',bold = True)
    style_paragraph_add(wenzi='本次小区级数据来源于双套系统摸底数据，十堰由于摸底存在问题，采用软采数据。')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='2019年第三季度软采数据评估情况：')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='全省分地市情况：')
    style_paragraph_add(wenzi='、'.join(dishi_fgl_qian3),bold = True)
    style_paragraph_add(wenzi='覆盖较好，')
    style_paragraph_add(wenzi='、'.join(dishi_fgl_hou3),bold = True)
    style_paragraph_add(wenzi='覆盖相对较差。')
    #插入表格并赋值
    data = data_fgl.copy()
    document.styles['Table Grid'].font.name = u'微软雅黑'
    document.styles['Table Grid']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = document.add_paragraph('')
 
    p = document.add_paragraph('')
    if data_fgl_qx_ok.shape[0]<3:
        style_paragraph_add(wenzi=dishi_name+'区县指标如下：')
    else:
        style_paragraph_add(wenzi=dishi_name+'各区县中，')
        style_paragraph_add(wenzi='、'.join(quxian_fgl_qian3),bold = True)
        style_paragraph_add(wenzi='覆盖较好，')
        style_paragraph_add(wenzi='、'.join(quxian_fgl_hou3),bold = True)
        style_paragraph_add(wenzi='覆盖相对较差。')
    #插入表格并赋值
    data = data_fgl_qx_ok.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #问题点统计---2.小区粒度
    def aggfunc_wenti(df):
        ruancai_count = df['eutrancell_cgi'].count()
        heidian_count = df['hd'].sum()
        ruofugai_count = df['弱覆盖'].sum()
        wenti_count = df['总问题小区'].sum()
        ser = pd.Series([ruancai_count, heidian_count,ruofugai_count, wenti_count,wenti_count/ruancai_count])
        return ser
    #地市问题点
    data_wenti = ruancai.groupby(by='region_name').apply(aggfunc_wenti)
    data_wenti.reset_index(inplace=True)
    data_wenti.columns = ['地市','总小区数','黑点小区数','弱覆盖小区数','问题小区数','问题小区占比']
    data_wenti.sort_values(by = '问题小区占比',ascending = True,inplace = True)
    #添加武汉中兴华为
    data_wenti_changjia = ruancai.loc[ruancai['region_name']=='武汉'].groupby(by='厂家名称').apply(aggfunc_wenti)
    data_wenti_changjia.reset_index(inplace=True)
    data_wenti_changjia.columns = ['地市','总小区数','黑点小区数','弱覆盖小区数','问题小区数','问题小区占比']
    data_wenti_changjia2 = data_wenti_changjia.loc[(data_wenti_changjia['地市']=='中兴') | (data_wenti_changjia['地市']=='华为')]
    data_wenti_changjia2.loc[data_wenti_changjia2['地市']=='中兴','地市']='武汉中兴'
    data_wenti_changjia2.loc[data_wenti_changjia2['地市']=='华为','地市']='武汉华为'
    data_wenti = pd.concat([data_wenti,data_wenti_changjia2])
    #添加全网指标
    data_wenti_quanwang = ruancai.groupby(by='全网').apply(aggfunc_wenti)
    data_wenti_quanwang.reset_index(inplace=True)
    data_wenti_quanwang.columns = ['地市','总小区数','黑点小区数','弱覆盖小区数','问题小区数','问题小区占比']
    data_wenti = pd.concat([data_wenti,data_wenti_quanwang])
    data_wenti['问题小区占比'] = data_wenti['问题小区占比'].apply(lambda x: format(x, '.2%'))    #Series.apply()让序列的值依次在lambda函数中执行； data['线损率']由小数转化为百分数
    
    dishi_wenti_qian3=list(data_wenti.iloc[0:3,0])
    dishi_wenti_hou3 = list(data_wenti.iloc[-6:-3,0])
    dishi_wenti_hou3.reverse()
    #武汉问题点
    data_wenti_qx = ruancai.loc[ruancai['对应栅格地']==dishi_name].groupby(by='对应栅格区县').apply(aggfunc_wenti)
    data_wenti_qx.reset_index(inplace=True)
    data_wenti_qx.columns = ['区县','总小区数','黑点小区数','弱覆盖小区数','问题小区数','问题小区占比']
    data_wenti_qx.sort_values(by = '问题小区占比',ascending = True,inplace = True)
    data_wenti_qx['问题小区占比'] = data_wenti_qx['问题小区占比'].apply(lambda x: format(x, '.2%'))   

    #武汉指标
    data_wenti_wuhan = data_wenti.loc[data_wenti['地市']==dishi_name]
    data_wenti_wuhan.columns = ['区县','总小区数','黑点小区数','弱覆盖小区数','问题小区数','问题小区占比']
    #区县总计
    data_wenti_qx_ok = pd.concat([data_wenti_qx,data_wenti_wuhan])

    if data_wenti_qx_ok.shape[0]>=7:
        quxian_wenti_qian3 = list(data_wenti_qx_ok.iloc[0:3,0])
        quxian_wenti_hou3 = list(data_wenti_qx_ok.iloc[-4:-1,0])
        quxian_wenti_hou3.reverse()
    elif data_wenti_qx_ok.shape[0]>=5:
        quxian_wenti_qian3 = list(data_wenti_qx_ok.iloc[0:2,0])
        quxian_wenti_hou3 = list(data_wenti_qx_ok.iloc[-3:-1,0])
        quxian_wenti_hou3.reverse()
    else:
        quxian_wenti_qian3 = list(data_wenti_qx_ok.iloc[0:1,0])
        quxian_wenti_hou3 = list(data_wenti_qx_ok.iloc[-2:-1,0])
        quxian_wenti_hou3.reverse()
    
    style_heading(wenzi ='2.小区粒度',jibie = 2,size = 12 , ziti=u'宋体',duiqi = 'LEFT')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='MR覆盖问题小区占比：',bold = True)
    style_paragraph_add(wenzi='全省分地市，武汉分区县MR覆盖问题小区占比情况')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='指标定义：',bold = True)
    style_paragraph_add(wenzi='MR覆盖问题小区占比=（小区级MR覆盖率<90%或MR竞对黑点有效小区）/总小区数')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='数据来源：',bold = True)
    style_paragraph_add(wenzi='本次小区级数据来源于双套系统摸底数据，十堰由于摸底存在问题，采用软采数据。')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='全省分地市情况：')
    style_paragraph_add(wenzi='、'.join(dishi_wenti_qian3),bold = True)
    style_paragraph_add(wenzi='覆盖较好，MR问题小区较少。')
    style_paragraph_add(wenzi='、'.join(dishi_wenti_hou3),bold = True)
    style_paragraph_add(wenzi='覆盖相对较差，MR问题小区较多。')
    
    #插入表格并赋值
    data = data_wenti.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    if data_wenti_qx_ok.shape[0]<3:
        style_paragraph_add(wenzi=dishi_name+'各区县指标如下：')
    else:
        style_paragraph_add(wenzi=dishi_name+'各区县中，')
        style_paragraph_add(wenzi='、'.join(quxian_wenti_qian3),bold = True)
        style_paragraph_add(wenzi='覆盖较好，MR问题小区较少。')
        style_paragraph_add(wenzi='、'.join(quxian_wenti_hou3),bold = True)
        style_paragraph_add(wenzi='覆盖相对较差，MR问题小区较多。')
    
    #插入表格并赋值
    data = data_wenti_qx_ok.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #3.大数据栅格粒度
    def aggfunc_grid(df):
        grid_count = df['a.grid_id'].count()
        grid_100_count = df['大于100采样点'].sum()
        df1 = df.loc[df['大于100采样点']==1]
        grid_fgl936_count = df1['弱覆盖栅格'].sum()
        ser = pd.Series([grid_count, grid_100_count,grid_fgl936_count, grid_fgl936_count/grid_100_count])
        return ser
    #地市问题点
    data_grid = data_merge_quchong.groupby(by='a.region_name').apply(aggfunc_grid)
    data_grid.reset_index(inplace=True)
    data_grid.columns = ['19年9月','总栅格','大于100采样点','弱覆盖栅格数量','大数据弱覆盖栅格占比']
    data_grid.sort_values(by = '大数据弱覆盖栅格占比',ascending = True,inplace = True)
    #添加全网指标
    data_grid_quanwang = data_merge_quchong.groupby(by='全网').apply(aggfunc_grid)
    data_grid_quanwang.reset_index(inplace=True)
    data_grid_quanwang.columns = ['19年9月','总栅格','大于100采样点','弱覆盖栅格数量','大数据弱覆盖栅格占比']
    data_grid = pd.concat([data_grid,data_grid_quanwang])
    data_grid['大数据弱覆盖栅格占比'] = data_grid['大数据弱覆盖栅格占比'].apply(lambda x: format(x, '.2%'))    #Series.apply()让序列的值依次在lambda函数中执行； data['线损率']由小数转化为百分数

    grid_zhanbi_qian3 = list(data_grid.iloc[0:3,0])
    grid_zhanbi_hou3 = list(data_grid.iloc[-4:-1,0])
    grid_zhanbi_hou3.reverse()
    #区县问题点
    data_grid_qx = data_merge_quchong.loc[data_merge_quchong['dishi_xxx']==dishi_name].groupby(by='区县').apply(aggfunc_grid)
    data_grid_qx.reset_index(inplace=True)
    data_grid_qx.columns = ['19年9月','总栅格','大于100采样点','弱覆盖栅格数量','大数据弱覆盖栅格占比']
    data_grid_qx.sort_values(by = '大数据弱覆盖栅格占比',ascending = True,inplace = True)
    data_grid_qx['大数据弱覆盖栅格占比'] = data_grid_qx['大数据弱覆盖栅格占比'].apply(lambda x: format(x, '.2%'))   

    #武汉指标
    data_grid_wuhan = data_grid.loc[data_grid['19年9月']==dishi_name]
    data_grid_wuhan.columns = ['19年9月','总栅格','大于100采样点','弱覆盖栅格数量','大数据弱覆盖栅格占比']
    data_grid_qx_ok = pd.concat([data_grid_qx,data_grid_wuhan])

    if data_grid_qx_ok.shape[0]>=7:
        grid_quxian_zhanbi_qian3 = list(data_grid_qx_ok.iloc[0:3,0])
        grid_quxian_zhanbi_hou3 = list(data_grid_qx_ok.iloc[-4:-1,0])
        grid_quxian_zhanbi_hou3.reverse()
    elif data_grid_qx_ok.shape[0]>=5:
        grid_quxian_zhanbi_qian3 = list(data_grid_qx_ok.iloc[0:2,0])
        grid_quxian_zhanbi_hou3 = list(data_grid_qx_ok.iloc[-3:-1,0])
        grid_quxian_zhanbi_hou3.reverse()
    else:
        grid_quxian_zhanbi_qian3 = list(data_grid_qx_ok.iloc[0:1,0])
        grid_quxian_zhanbi_hou3 = list(data_grid_qx_ok.iloc[-2:-1,0])
        grid_quxian_zhanbi_hou3.reverse()
        
    style_heading(wenzi ='3.大数据栅格粒度',jibie = 2,size = 12 , ziti=u'宋体',duiqi = 'LEFT')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='1)大数据MR问题栅格占比:',bold = True)
    style_paragraph_add(wenzi='全省分地市，武汉分区县大数据问题栅格占比情况')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='指标定义：',bold = True)
    style_paragraph_add(wenzi='大数据MR问题栅格占比=大数据问题栅格数/总有效栅格数')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='大数据问题栅格：有效栅格MR覆盖率小于93.6%(此处门限定义为目标网设定全网达到96%的换算值)')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='总有效栅格数：栅格采样点数大于100')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='全省分地市情况：')
    style_paragraph_add(wenzi='、'.join(grid_zhanbi_qian3),bold = True)
    style_paragraph_add(wenzi='覆盖较好，MR弱覆盖栅格较少。')
    style_paragraph_add(wenzi='、'.join(grid_zhanbi_hou3),bold = True)
    style_paragraph_add(wenzi='覆盖相对较差，MR弱覆盖栅格较多。')
    p = document.add_paragraph('')
    style_paragraph_hongse_add(wenzi='注：大数据问题栅格标准是栅格覆盖率小于93.6%，故此处指标较高。')
    
    #插入表格并赋值
    data = data_grid.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = document.add_paragraph('')

    p = document.add_paragraph('')
    if data_grid_qx_ok.shape[0]<3:
        style_paragraph_add(wenzi=dishi_name+'各区县指标如下：')
    else:
        style_paragraph_add(wenzi=dishi_name+'各区县中，')
        style_paragraph_add(wenzi='、'.join(grid_quxian_zhanbi_qian3),bold = True)
        style_paragraph_add(wenzi='覆盖较好，MR弱覆盖栅格较少。')
        style_paragraph_add(wenzi='、'.join(grid_quxian_zhanbi_hou3),bold = True)
        style_paragraph_add(wenzi='覆盖相对较差，MR弱覆盖栅格较多。')

    #插入表格并赋值
    data = data_grid_qx_ok.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
    def aggfunc_grid_yanzhong(df):
        grid_count = df['a.grid_id'].count()
        grid_100_count = df['大于100采样点'].sum()
        df1 = df.loc[df['大于100采样点']==1]
        grid_fgl936_count = df1['严重弱覆盖栅格'].sum()
        ser = pd.Series([grid_count, grid_100_count,grid_fgl936_count, grid_fgl936_count/grid_100_count])
        return ser

    #地市问题点
    
    data_grid_yz = data_merge_quchong.groupby(by='a.region_name').apply(aggfunc_grid_yanzhong)
    data_grid_yz.reset_index(inplace=True)
    data_grid_yz.columns = ['19年9月','总栅格','大于100采样点','严重弱覆盖栅格数量','严重弱覆盖栅格占比']
    data_grid_yz.sort_values(by = '严重弱覆盖栅格占比',ascending = True,inplace = True)

    #添加全网指标
    
    data_grid_quanwang_yz = data_merge_quchong.groupby(by='全网').apply(aggfunc_grid_yanzhong)
    data_grid_quanwang_yz.reset_index(inplace=True)
    data_grid_quanwang_yz.columns = ['19年9月','总栅格','大于100采样点','严重弱覆盖栅格数量','严重弱覆盖栅格占比']
    data_grid_yz = pd.concat([data_grid_yz,data_grid_quanwang_yz])
    data_grid_yz['严重弱覆盖栅格占比'] = data_grid_yz['严重弱覆盖栅格占比'].apply(lambda x: format(x, '.2%'))    #Series.apply()让序列的值依次在lambda函数中执行； data['线损率']由小数转化为百分数
    data_grid_yz
    
    grid_yanzhong_qian3 = list(data_grid_yz.iloc[0:3,0])
    grid_yanzhong_hou3 = list(data_grid_yz.iloc[-4:-1,0])
    grid_yanzhong_hou3.reverse()

    #区县问题点
    
    data_grid_qx_yz = data_merge_quchong.loc[data_merge_quchong['dishi_xxx']==dishi_name].groupby(by='区县').apply(aggfunc_grid_yanzhong)
    data_grid_qx_yz.reset_index(inplace=True)
    data_grid_qx_yz.columns = ['19年9月','总栅格','大于100采样点','严重弱覆盖栅格数量','严重弱覆盖栅格占比']
    data_grid_qx_yz.sort_values(by = '严重弱覆盖栅格占比',ascending = True,inplace = True)
    data_grid_qx_yz['严重弱覆盖栅格占比'] = data_grid_qx_yz['严重弱覆盖栅格占比'].apply(lambda x: format(x, '.2%')) 
    
    #武汉指标
    data_grid_wuhan_yz = data_grid_yz.loc[data_grid_yz['19年9月']==dishi_name]
    data_grid_wuhan_yz.columns = ['19年9月','总栅格','大于100采样点','严重弱覆盖栅格数量','严重弱覆盖栅格占比']
    data_grid_qx_yz_ok = pd.concat([data_grid_qx_yz,data_grid_wuhan_yz])
    data_grid_qx_yz_ok


    if data_grid_qx_yz_ok.shape[0]>=7:
        grid_yanzhong_quxian_qian3 = list(data_grid_qx_yz_ok.iloc[0:3,0])
        grid_yanzhong_quxian_hou3 = list(data_grid_qx_yz_ok.iloc[-4:-1,0])
        grid_yanzhong_quxian_hou3.reverse()
    elif data_grid_qx_yz_ok.shape[0]>=5:
        grid_yanzhong_quxian_qian3 = list(data_grid_qx_yz_ok.iloc[0:2,0])
        grid_yanzhong_quxian_hou3 = list(data_grid_qx_yz_ok.iloc[-3:-1,0])
        grid_yanzhong_quxian_hou3.reverse()
    else:
        grid_yanzhong_quxian_qian3 = list(data_grid_qx_yz_ok.iloc[0:1,0])
        grid_yanzhong_quxian_hou3 = list(data_grid_qx_yz_ok.iloc[-2:-1,0])
        grid_yanzhong_quxian_hou3.reverse()

    p = document.add_paragraph('')
    style_paragraph_add(wenzi='\n2)严重弱覆盖栅格占比：',bold = True)
    style_paragraph_add(wenzi='全省分地市，武汉分区县大数据问题栅格占比情况')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='指标定义：',bold = True)
    style_paragraph_add(wenzi='大数据严重弱覆盖栅格占比=大数据严重问题栅格数/总有效栅格数大数据严重问题栅格：有效栅格MR覆盖率小于70%.')
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='全省分地市情况：')
    style_paragraph_add(wenzi='、'.join(grid_yanzhong_qian3),bold = True)
    style_paragraph_add(wenzi='覆盖较好，MR弱覆盖栅格较少。')
    style_paragraph_add(wenzi='、'.join(grid_yanzhong_hou3),bold = True)
    style_paragraph_add(wenzi='覆盖相对较差，MR弱覆盖栅格较多。')

    #插入表格并赋值
    data = data_grid_yz.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
       
    p = document.add_paragraph('')
    style_paragraph_add('\n') 
    p = document.add_paragraph('')
    if data_grid_qx_yz_ok.shape[0]<3:
        style_paragraph_add(wenzi=dishi_name+'各区县指标如下：')
    else:
        style_paragraph_add(wenzi=dishi_name+'各区县中，')
        style_paragraph_add(wenzi='、'.join(grid_yanzhong_quxian_qian3),bold = True)
        style_paragraph_add(wenzi='覆盖较好，MR严重弱覆盖栅格较少。')
        style_paragraph_add(wenzi='、'.join(grid_yanzhong_quxian_hou3),bold = True)
        style_paragraph_add(wenzi='覆盖相对较差，MR严重弱覆盖栅格较多。')
    
    #插入表格并赋值
    data = data_grid_qx_yz_ok.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            
    #添加word内容=======================
    style_heading(wenzi ='二、	整体评估',jibie = 2,size = 14 , ziti=u'宋体',duiqi = 'LEFT')
    style_paragraph(wenzi = '各指标项分值见分值列：地市指标按指标情况线性取值。')
    
    #插入表格并赋值
    data = biaozhun.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #得分计算——地市
    list_dishi2 = ['孝感','襄阳','咸宁','随州','十堰','荆州','荆门','黄石','鄂州','黄冈','恩施','宜昌','潜江','天门','江汉','武汉','全网']
    pingfen_dishi = pd.DataFrame({'地市': list_dishi2})
    pingfen_dishi['全网MR覆盖率']=0
    pingfen_dishi['得分1']=0
    pingfen_dishi['问题小区占比']=0
    pingfen_dishi['得分2']=0
    pingfen_dishi['大数据MR问题栅格占比']=0
    pingfen_dishi['得分3']=0
    pingfen_dishi['严重弱覆盖栅格占比']=0
    pingfen_dishi['得分4']=0
    pingfen_dishi['总得分']=0
    
    pingfen_dishi.set_index('地市',inplace = True)
    data_fgl_index = data_fgl.set_index('全网-指标')
    data_wenti_index = data_wenti.set_index('地市')
    data_grid_index = data_grid.set_index('19年9月')
    data_grid_yz_index = data_grid_yz.set_index('19年9月')
    pingfen_dishi['全网MR覆盖率'].update(data_fgl_index['全网移动覆盖率'])
    pingfen_dishi['问题小区占比'].update(data_wenti_index['问题小区占比'])
    pingfen_dishi['大数据MR问题栅格占比'].update(data_grid_index['大数据弱覆盖栅格占比'])
    pingfen_dishi['严重弱覆盖栅格占比'].update(data_grid_yz_index['严重弱覆盖栅格占比'])
    pingfen_dishi['得分1'] = pd.to_numeric(pingfen_dishi['全网MR覆盖率'].str.replace('%',''))*20/100
    pingfen_dishi['得分2'] = 10-pd.to_numeric(pingfen_dishi['问题小区占比'].str.replace('%',''))*10/100
    pingfen_dishi['得分3'] = 10-pd.to_numeric(pingfen_dishi['大数据MR问题栅格占比'].str.replace('%',''))*10/100
    pingfen_dishi['得分4'] = 60-pd.to_numeric(pingfen_dishi['严重弱覆盖栅格占比'].str.replace('%',''))/100*60
    pingfen_dishi['总得分'] = pingfen_dishi['得分1']+pingfen_dishi['得分2']+pingfen_dishi['得分3']+pingfen_dishi['得分4']
    pingfen_dishi_1 = pingfen_dishi.round(decimals = 2)
    pingfen_dishi_1.sort_values('总得分',ascending=False,inplace = True)
    pingfen_dishi_1.reset_index(inplace = True)
    
    pingfen_dishi_res = pd.concat([pingfen_dishi_1.loc[pingfen_dishi_1['地市']!='全网'],pingfen_dishi_1.loc[pingfen_dishi_1['地市']=='全网']])
    pingfen_dishi_res.reset_index(drop=True,inplace = True)
    
    pingfen_dishi_res.shape
    if data_grid_qx_yz_ok.shape[0]>=7:
        pingfen_dishi_qian3 = list(pingfen_dishi_res.iloc[0:3,0])
        pingfen_dishi_hou3 = list(pingfen_dishi_res.iloc[-4:-1,0])
        pingfen_dishi_hou3.reverse()
    elif pingfen_dishi_res.shape[0]>=5:
        pingfen_dishi_qian3 = list(pingfen_dishi_res.iloc[0:2,0])
        pingfen_dishi_hou3 = list(pingfen_dishi_res.iloc[-3:-1,0])
        pingfen_dishi_hou3.reverse()
    else:
        pingfen_dishi_qian3 = list(pingfen_dishi_res.iloc[0:1,0])
        pingfen_dishi_hou3 = list(pingfen_dishi_res.iloc[-2:-1,0])
        pingfen_dishi_hou3.reverse()
    wuhan_paiming = pingfen_dishi_res[pingfen_dishi_res['地市'].isin(['武汉'])].index[0]+1

    p = document.add_paragraph('')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='最后依据地市每项得分的总和得到地市总得分，排序情况如下：')
    style_paragraph_add(wenzi='、'.join(pingfen_dishi_qian3),bold = True)
    style_paragraph_add(wenzi='得分最高。')
    style_paragraph_add(wenzi='、'.join(pingfen_dishi_hou3),bold = True)
    style_paragraph_add(wenzi='得分最低。')
    style_paragraph_hongse_add(wenzi='武分排名'+str(wuhan_paiming)+'。')
    #插入表格并赋值
    data = pingfen_dishi_res.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #得分-区县
    pingfen_quxian = data_fgl_qx_ok[['区域','全网移动覆盖率']]
    pingfen_quxian = pingfen_quxian.reindex(columns=['区域','全网移动覆盖率','得分1','问题小区占比','得分2','大数据MR问题栅格占比','得分3','严重弱覆盖栅格占比','得分4','总得分'])
    pingfen_quxian1 = pingfen_quxian.set_index('区域')
    data_wenti_qx_ok_index = data_wenti_qx_ok.set_index('区县')
    data_grid_qx_ok_index = data_grid_qx_ok.set_index('19年9月')
    data_grid_qx_yz_ok_index = data_grid_qx_yz_ok.set_index('19年9月')
    pingfen_quxian1['问题小区占比'].update(data_wenti_qx_ok_index['问题小区占比'])
    pingfen_quxian1['大数据MR问题栅格占比'].update(data_grid_qx_ok_index['大数据弱覆盖栅格占比'])
    pingfen_quxian1['严重弱覆盖栅格占比'].update(data_grid_qx_yz_ok_index['严重弱覆盖栅格占比'])
    pingfen_quxian1['得分1'] = pd.to_numeric(pingfen_quxian1['全网移动覆盖率'].str.replace('%',''))*20/100
    pingfen_quxian1['得分2'] = 10-pd.to_numeric(pingfen_quxian1['问题小区占比'].str.replace('%',''))*10/100
    pingfen_quxian1['得分3'] = 10-pd.to_numeric(pingfen_quxian1['大数据MR问题栅格占比'].str.replace('%',''))*10/100
    pingfen_quxian1['得分4'] = 60-pd.to_numeric(pingfen_quxian1['严重弱覆盖栅格占比'].str.replace('%',''))/100*60
    pingfen_quxian1['总得分'] = pingfen_quxian1['得分1']+pingfen_quxian1['得分2']+pingfen_quxian1['得分3']+pingfen_quxian1['得分4']
    pingfen_quxian_1 = pingfen_quxian1.round(decimals = 2)
    pingfen_quxian_1.sort_values('总得分',ascending=False,inplace = True)
    pingfen_quxian_1.reset_index(inplace = True)
    
    pingfen_quxian_res = pd.concat([pingfen_quxian_1.loc[pingfen_quxian_1['区域']!=dishi_name],pingfen_quxian_1.loc[pingfen_quxian_1['区域']==dishi_name]])
    pingfen_quxian_res.reset_index(drop=True,inplace = True)
    
    if pingfen_quxian_res.shape[0]>=7:
        pingfen_quxian_qian3 = list(pingfen_quxian_res.iloc[0:3,0])
        pingfen_quxian_hou3 = list(pingfen_quxian_res.iloc[-4:-1,0])
        pingfen_quxian_hou3.reverse()
    elif pingfen_quxian_res.shape[0]>=5:
        pingfen_quxian_qian3 = list(pingfen_quxian_res.iloc[0:2,0])
        pingfen_quxian_hou3 = list(pingfen_quxian_res.iloc[-3:-1,0])
        pingfen_quxian_hou3.reverse()
    else:
        pingfen_quxian_qian3 = list(pingfen_quxian_res.iloc[0:1,0])
        pingfen_quxian_hou3 = list(pingfen_quxian_res.iloc[-2:-1,0])
        pingfen_quxian_hou3.reverse()

    p = document.add_paragraph('')
    p = document.add_paragraph('')
    style_paragraph_add(wenzi='依据'+dishi_name+'区县的每项得分总和得到各区县的总得分，排序如下：')
    if pingfen_quxian_res.shape[0]<3:
        pass
    else:
        style_paragraph_add(wenzi='、'.join(pingfen_quxian_qian3),bold = True)
        style_paragraph_add(wenzi='得分最高。')
        style_paragraph_add(wenzi='、'.join(pingfen_quxian_hou3),bold = True)
        style_paragraph_add(wenzi='得分最低。')
    #插入表格并赋值
    data = pingfen_quxian_res.copy()
    table = document.add_table(rows = data.shape[0]+1 , cols = data.shape[1],style = 'Table Grid')
    table.autofit=True
    table.rows[0].height=Cm(1)
    # 列名赋值
    hdr_cells = table.rows[0].cells
    for row in list(data.columns):
        for num in range(0, data.shape[1]):
            hdr_cells[num].text = u'' + list(data.columns)[num]
            table.cell(0,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#BFBFBF"/>'.format(nsdecls('w')))
            table.cell(0,num)._tc.get_or_add_tcPr().append(shading_elm_2)
        break
    # 内容赋值
    for row in range(1,data.shape[0]+1):
        row_cells = table.rows[row].cells
        for num in range(0, data.shape[1]):
            row_cells[num].text = u'' + str(data.iloc[row-1,num])
            table.cell(row,num).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            table.cell(row,num).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.save(dishi_name+'三季度评估报告.docx')  #保存文档
for name in list_dishi:
    print(name,'开始')
    to_word(dishi_name = name)
    print(name,'完成')